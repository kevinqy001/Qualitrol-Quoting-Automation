"""One-off builder: create the blank Step 4 quotation template.

Derives ``Quotation_Template.docx`` (saved next to this script) from the
official filled sample quotation by neutralising ALL project-specific content
while preserving the full structure, styling, page header / logo and the legal
Terms & Conditions verbatim. Insights from ``Gemba Samples/2/Templetes/
MEA Template 2026.docx`` (the official blank master) were used to decide what is
project-specific vs. standard boilerplate.

What gets neutralised / blanked (so the template is fully reusable):
  * Page-1 sales contact block  : sales name + personal email -> placeholders
                                  (filled per run via Step 4 --sales-name/-email),
                                  and the mailto hyperlink target is genericised.
  * Page header                 : QUOTE/SFDC # -> "XXXX-0".
  * Revision + Project Info      : data rows / value cells blanked.
  * Customer Requirement / Notes : reduced to one blank item (filled per run).
  * Equipment schedule (1a)      : tables emptied to the QTY|DESCRIPTION header.
  * Page-2 Table of Contents     : cached project headings scrubbed AND the doc
                                   is set to update fields (TOC + date) on open.
  * Project tokens everywhere    : IBRI / OMAN / UAE 2026 / Ibri doc names, etc.
  * All prices                   : "$ ______".
Kept verbatim: legal T&C (Schedules A/B/C), Qualitrol factory ("Belfast"),
company emails (quotes@ / belfastorders@), Qualitrol logo & styles.

The blank template keeps the SAME table count and positions as the source so
``quotation_docgen`` (which relies on fixed table indices) behaves identically.

Run once (or re-run if the source layout changes):
    python "Step 4 _ Generate Quotation/build_blank_template.py"
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

_REPO_ROOT = Path(__file__).resolve().parent.parent
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

SOURCE = (_REPO_ROOT / "Gemba Samples" / "1" / "1" / "773306" / "3. QUOTE"
          / "108704-749714.docx")
OUTPUT = Path(__file__).resolve().parent / "Quotation_Template.docx"

_MONEY_RE = re.compile(r"(\$\s*)[\d][\d,]*(?:\.\d+)?")

# Placeholders that Step 4 fills per run (see quotation_docgen). They appear on
# the page-1 cover content control and elsewhere.
SALES_NAME_PLACEHOLDER = "[Sales Contact Name]"
SALES_EMAIL_PLACEHOLDER = "sales.contact@qualitrolcorp.com"
CUSTOMER_PLACEHOLDER = "[CUSTOMER]"
QUOTE_NUMBER_PLACEHOLDER = "XXXX-0"

# Project-specific text -> generic placeholder. Applied to EVERY <w:t> in the
# body (catches content controls, table cells, text boxes and TOC cache), so it
# is order-independent. Longest/most-specific first.
_NEUTRALISE = [
    ("Aditya Taneja", SALES_NAME_PLACEHOLDER),
    ("aditya.taneja@qualitrolcorp.com", SALES_EMAIL_PLACEHOLDER),
    # Cover / header quote number (handle whole and split-run forms).
    ("108704-749714", QUOTE_NUMBER_PLACEHOLDER),
    ("749714", "0"),
    ("108704", "XXXX"),
    # Customer name (page-1 cover + body).
    ("GCCIA", CUSTOMER_PLACEHOLDER),
    ("gccia", CUSTOMER_PLACEHOLDER),
    # Site / location.
    ("IBRI SS ", ""),
    ("IBRI 400kV", "400kV"),
    ("IBRI 33kV", "33kV"),
    ("IBRI ", ""),
    ("IBRI", ""),
    ("CIP OMAN", "CIP [Destination]"),
    ("Customer Site UAE 2026", "Customer Site [Location] [Year]"),
    ("UAE 2026", "[Location] [Year]"),
    ("UAE", "[Location]"),
    ("Qualitrol Factory 2026", "Qualitrol Factory [Year]"),
]

_HEADING_STYLES = {"Title", "Heading 1", "Heading 2", "Heading Level 3"}


def _set_text_keep_format(paragraph, text: str) -> None:
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _set_cell_text(cell, text: str) -> None:
    para = cell.paragraphs[0]
    _set_text_keep_format(para, text)
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)


def _neutralise_all_text(doc) -> None:
    """Replace project-specific tokens in every <w:t> across the whole body."""
    for t in doc.element.body.iter(qn("w:t")):
        if not t.text:
            continue
        new = t.text
        for old, repl in _NEUTRALISE:
            if old in new:
                new = new.replace(old, repl)
        if new != t.text:
            t.text = new


def _fix_split_quote_number(doc) -> None:
    """Scrub the cover quote/SFDC number even when split across runs.

    Word stores "108704-749714" as several tiny runs (e.g. "10","8704","-",
    "749714"), so per-run replace cannot match it. For any paragraph whose
    *joined* text still carries the old number, merge the runs' text into the
    first run, apply the neutralisation, and blank the rest.
    """
    for p in doc.element.body.iter(qn("w:p")):
        ts = list(p.iter(qn("w:t")))
        joined = "".join(t.text or "" for t in ts)
        if "108704" not in joined and "749714" not in joined:
            continue
        new = joined
        for old, repl in _NEUTRALISE:
            if old in new:
                new = new.replace(old, repl)
        if ts:
            ts[0].text = new
            for t in ts[1:]:
                t.text = ""


def _genericise_mailto(doc) -> None:
    """Repoint any personal sales mailto link to the generic quotes inbox."""
    for rel in list(doc.part.rels.values()):
        try:
            target = rel.target_ref or ""
        except Exception:  # noqa: BLE001
            continue
        if "mailto:" in target and "aditya" in target.lower():
            rel._target = f"mailto:{SALES_EMAIL_PLACEHOLDER}"


def _update_fields_on_open(doc) -> None:
    """Tell Word to refresh fields (TOC + date) when the document is opened."""
    settings = doc.settings.element
    el = settings.find(qn("w:updateFields"))
    if el is None:
        el = settings.makeelement(qn("w:updateFields"), {qn("w:val"): "true"})
        settings.insert(0, el)
    else:
        el.set(qn("w:val"), "true")


def _blank_list_after(doc, predicate) -> None:
    paras = doc.paragraphs
    anchor = next((p for p in paras if predicate(p)), None)
    if anchor is None:
        return
    items, started = [], False
    for p in paras:
        if p._p is anchor._p:
            started = True
            continue
        if started:
            if p.style.name in _HEADING_STYLES:
                break
            items.append(p)
    if not items:
        return
    _set_text_keep_format(items[0], "")
    for extra in items[1:]:
        extra._p.getparent().remove(extra._p)


def build() -> Path:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Source quotation not found: {SOURCE}")
    doc = Document(str(SOURCE))

    # 1. Header -> placeholder SFDC number.
    for section in doc.sections:
        for para in section.header.paragraphs:
            if "QUOTE/SFDC" in para.text:
                _set_text_keep_format(para, "QUOTE/SFDC # XXXX-0")

    tables = doc.tables

    # 2. Revision table: keep header row, blank the data rows.
    if tables:
        for row in tables[0].rows[1:]:
            for cell in row.cells:
                _set_cell_text(cell, "")

    # 3. Project Information table: blank the value column.
    if len(tables) > 1:
        for row in tables[1].rows:
            _set_cell_text(row.cells[1], "")

    # 4. Equipment schedule tables (Attachment 1a, indices 2..8): keep the
    #    QTY | PRODUCT DESCRIPTION header row, blank the rest. Tables are kept
    #    in place so docgen's fixed-index logic still works.
    for tbl in tables[2:9]:
        for row in list(tbl.rows)[1:]:
            row._tr.getparent().remove(row._tr)

    # 5. Blank prices in body paragraphs.
    for para in doc.paragraphs:
        if "$" in para.text:
            _set_text_keep_format(
                para, _MONEY_RE.sub(lambda m: f"{m.group(1)}______", para.text)
            )

    # 6. Customer Requirement + Important Notes: one blank placeholder item.
    _blank_list_after(doc, lambda p: "based our quotation" in p.text.lower())
    _blank_list_after(doc, lambda p: p.text.strip().lower().startswith("important notes"))

    # 7. Blank prices inside table cells (service / training notes etc.).
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "$" in para.text:
                        _set_text_keep_format(
                            para, _MONEY_RE.sub(lambda m: f"{m.group(1)}______", para.text)
                        )

    # 8. Neutralise project-specific tokens + sales block across the WHOLE body
    #    (content controls, table cells, text boxes, TOC cache).
    _neutralise_all_text(doc)
    _fix_split_quote_number(doc)

    # 9. Repoint the personal sales mailto link to a generic inbox.
    _genericise_mailto(doc)

    # 10. Refresh TOC + date fields when the document is opened in Word.
    _update_fields_on_open(doc)

    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    out = build()
    print(f"Blank template written to: {out}")

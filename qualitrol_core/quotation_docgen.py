"""Step 4 - Standard Qualitrol Quotation document builder.

STANDARD (must be honoured by Step 4 now and in future):
    The quotation MUST match the official Qualitrol template
    ``Gemba Samples/1/1/773306/3. QUOTE/108704-749714.docx`` exactly — every
    page, every section, every style/format and the full legal Terms &
    Conditions. We therefore *clone the template file itself* and fill only the
    dynamic regions in place, so all styling, headers, the Qualitrol logo and
    the boilerplate carry over verbatim.

    Filling rules:
      * Where the run already has content (project info, source documents,
        scope notes, equipment/BOQ, application scenarios) -> fill it in.
      * Where content is not available yet (all pricing — pending the Step 3
        pricing layer) -> leave it BLANK ("$ ______") rather than inventing it.

Dynamic regions filled from the Step 1 + Step 2 outputs:
    - Page header  : QUOTE/SFDC number
    - Revision table
    - Project Information table (customer / tender ref / project / location / SFDC)
    - Customer Requirement (the documents the quote is based on)   <- Step 1
    - Important Notes on Offer (scope assumptions)                 <- Step 2
    - Attachment 1a equipment schedule (the matched BOQ)           <- Step 2
    - All monetary values blanked (pending Step 3 pricing)
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from typing import Optional

# Paragraph styles that mark the start of a new section/block.
_HEADING_STYLES = {"Title", "Heading 1", "Heading 2", "Heading Level 3"}

# Matches a money amount after a "$" so we can blank the number but keep any
# trailing unit (e.g. "/day", "/visit", " per week").
_MONEY_RE = re.compile(r"(\$\s*)[\d][\d,]*(?:\.\d+)?")


class QuotationMeta:
    """Editable header metadata for the quotation.

    Anything not provided keeps the template's placeholder so the document is
    obviously a draft pending sales review.
    """

    def __init__(
        self,
        project_id: str,
        quote_number: Optional[str] = None,
        customer: Optional[str] = None,
        project_name: Optional[str] = None,
        location: Optional[str] = None,
        tender_ref: Optional[str] = None,
        sfdc_number: Optional[str] = None,
        currency: str = "USD",
        validity_days: int = 90,
        sales_name: Optional[str] = None,
        sales_email: Optional[str] = None,
    ) -> None:
        self.project_id = project_id
        self.quote_number = quote_number or sfdc_number or project_id
        self.customer = customer
        self.project_name = project_name
        self.location = location
        self.tender_ref = tender_ref
        self.sfdc_number = sfdc_number or project_id
        self.currency = currency
        self.validity_days = validity_days
        self.sales_name = sales_name
        self.sales_email = sales_email


# Placeholders left in the blank template by build_blank_template.py.
SALES_NAME_PLACEHOLDER = "[Sales Contact Name]"
SALES_EMAIL_PLACEHOLDER = "sales.contact@qualitrolcorp.com"
CUSTOMER_PLACEHOLDER = "[CUSTOMER]"
QUOTE_NUMBER_PLACEHOLDER = "XXXX-0"


# --------------------------------------------------------------------------- #
# Low-level docx helpers
# --------------------------------------------------------------------------- #
def _set_text_keep_format(paragraph, text: str) -> None:
    """Set a paragraph's text while preserving the first run's formatting."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _set_cell_text(cell, text: str) -> None:
    """Replace a table cell's text, keeping the existing run formatting."""
    para = cell.paragraphs[0]
    _set_text_keep_format(para, text)
    # Drop any extra paragraphs in the cell.
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)


def _blank_money_in_text(text: str) -> str:
    return _MONEY_RE.sub(lambda m: f"{m.group(1)}______", text)


def _para_text(paragraph) -> str:
    return paragraph.text.strip()


def _find_para(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p):
            return p
    return None


def _collect_following_items(doc, anchor_para):
    """Paragraphs after ``anchor_para`` up to the next heading-styled paragraph.

    Compares underlying XML elements (``._p``) because python-docx returns fresh
    Paragraph wrapper objects on each ``doc.paragraphs`` access.
    """
    paras = doc.paragraphs
    start = None
    for idx, p in enumerate(paras):
        if p._p is anchor_para._p:
            start = idx
            break
    if start is None:
        return []
    items = []
    for p in paras[start + 1:]:
        if p.style.name in _HEADING_STYLES:
            break
        items.append(p)
    return items


def _replace_list_items(doc, items, new_texts: list[str]) -> None:
    """Rewrite a block of list paragraphs to ``new_texts`` (clone/remove to fit)."""
    if not items:
        return
    if not new_texts:
        new_texts = [""]
    # Reuse / add.
    last = items[-1]
    for i, text in enumerate(new_texts):
        if i < len(items):
            _set_text_keep_format(items[i], text)
        else:
            new_el = deepcopy(last._p)
            last._p.addnext(new_el)
            from docx.text.paragraph import Paragraph
            new_p = Paragraph(new_el, last._parent)
            _set_text_keep_format(new_p, text)
            last = new_p
    # Remove leftover original items.
    for extra in items[len(new_texts):]:
        extra._p.getparent().remove(extra._p)


def _remove_element(el) -> None:
    el.getparent().remove(el)


# --------------------------------------------------------------------------- #
# Section fillers
# --------------------------------------------------------------------------- #
def _fill_header(doc, meta: QuotationMeta) -> None:
    """Replace 'QUOTE/SFDC # <num>' in every section header."""
    for section in doc.sections:
        for para in section.header.paragraphs:
            if "QUOTE/SFDC" in para.text:
                new = re.sub(r"(QUOTE/SFDC\s*#\s*).*", rf"\g<1>{meta.quote_number}",
                             para.text)
                _set_text_keep_format(para, new)


def _fill_revision(doc) -> None:
    """First data row -> DRAFT; clear the rest."""
    if not doc.tables:
        return
    rev = doc.tables[0]
    if len(rev.rows) < 2:
        return
    from datetime import datetime
    row1 = rev.rows[1].cells
    _set_cell_text(row1[0], "DRAFT")
    _set_cell_text(row1[1], "Automated first-pass quotation (AI draft)")
    _set_cell_text(row1[2], datetime.now().strftime("%d/%m/%Y"))
    _set_cell_text(row1[3], "AI")
    for r in rev.rows[2:]:
        for c in r.cells:
            _set_cell_text(c, "")


def _fill_project_info(doc, meta: QuotationMeta) -> None:
    """Project Information table (4x2): fill the value column where we have data."""
    if len(doc.tables) < 2:
        return
    info = doc.tables[1]
    mapping = {
        "end user": meta.customer,
        "tender": meta.tender_ref,
        "project name": (" ".join(x for x in (meta.project_name, meta.location) if x)
                         or None),
        "sfdc": meta.sfdc_number,
    }
    for row in info.rows:
        label = row.cells[0].text.strip().lower()
        for key, value in mapping.items():
            if key in label and value:
                _set_cell_text(row.cells[1], value)
                break


def _fill_customer_requirement(doc, step1: dict) -> None:
    anchor = _find_para(doc, lambda p: "based our quotation" in p.text.lower())
    if anchor is None:
        return
    items = _collect_following_items(doc, anchor)
    docs = step1.get("documents", [])
    names = [f"{d.get('file_name', '')}" for d in docs] or ["[No source documents recorded]"]
    _replace_list_items(doc, items, names)


def _fill_important_notes(doc, step2: dict) -> None:
    anchor = _find_para(doc, lambda p: p.text.strip().lower().startswith("important notes")
                        and p.style.name in _HEADING_STYLES)
    if anchor is None:
        return
    items = _collect_following_items(doc, anchor)
    seen: set[str] = set()
    notes: list[str] = []
    for line in step2.get("draft_boq", []):
        a = (line.get("assumption") or "").strip()
        if a and a.lower() not in seen:
            seen.add(a.lower())
            notes.append(a)
    notes.append("Items flagged 'Needs Review' require confirmation before the "
                 "configuration and price are finalised (see Open Clarification "
                 "Questions / Missing Information).")
    _replace_list_items(doc, items, notes)


def _boq_description(line: dict) -> str:
    """Compose an equipment-schedule description from a BOQ line."""
    model = (line.get("product_model") or "").strip()
    desc = (line.get("product_description") or "").strip()
    head = f"{model} – {desc}" if model else desc
    status = (line.get("review_status") or "").strip()
    if status and status.lower() != "draft":
        head = f"{head}  [{status}]"
    return head


def _fill_equipment_schedule(doc, step1: dict, step2: dict) -> None:
    """Rebuild Attachment 1a as a single equipment table from the matched BOQ.

    The template's per-panel breakdown is project-specific; we keep the exact
    table format (QTY | PRODUCT DESCRIPTION, same style) but populate it with the
    actual matched BOQ. Remaining template panel tables are removed.
    """
    boq = step2.get("draft_boq", [])
    # Equipment tables are tables[2:9] in the template (panel schedules). The
    # first one is reused; the rest are removed along with their heading.
    tables = doc.tables
    if len(tables) < 3:
        return

    # Retitle the first panel heading -> generic "Main Scope Equipment Schedule".
    first_tbl = tables[2]
    prev = first_tbl._tbl.getprevious()
    from docx.oxml.ns import qn
    if prev is not None and prev.tag == qn("w:p"):
        from docx.text.paragraph import Paragraph
        _set_text_keep_format(Paragraph(prev, first_tbl._parent),
                              "Main Scope Equipment Schedule")

    # Rebuild rows of the first equipment table from the BOQ.
    for row in list(first_tbl.rows):
        row._tr.getparent().remove(row._tr)
    hdr = first_tbl.add_row().cells
    _set_cell_text(hdr[0], "QTY")
    _set_cell_text(hdr[1], "PRODUCT DESCRIPTION")
    for run in hdr[0].paragraphs[0].runs:
        run.bold = True
    for run in hdr[1].paragraphs[0].runs:
        run.bold = True
    for line in boq:
        cells = first_tbl.add_row().cells
        qty = line.get("quantity") or 0
        try:
            qty_txt = str(int(float(qty))) if float(qty) > 0 else ""
        except (TypeError, ValueError):
            qty_txt = ""
        _set_cell_text(cells[0], qty_txt)
        _set_cell_text(cells[1], _boq_description(line))

    # Remove the other panel headings + tables (tables index 3..8 = next 6).
    for tbl in tables[3:9]:
        prev = tbl._tbl.getprevious()
        if prev is not None and prev.tag == qn("w:p"):
            _remove_element(prev)
        _remove_element(tbl._tbl)


def _replace_body_text(doc, repl: dict[str, str]) -> None:
    """Replace placeholder substrings in every body <w:t> (incl. cover SDT)."""
    from docx.oxml.ns import qn

    repl = {k: v for k, v in repl.items() if v}
    if not repl:
        return
    for t in doc.element.body.iter(qn("w:t")):
        if not t.text:
            continue
        new = t.text
        for old, value in repl.items():
            if old in new:
                new = new.replace(old, value)
        if new != t.text:
            t.text = new


def _fill_cover(doc, meta: QuotationMeta) -> None:
    """Fill the page-1 cover content control (customer + quote/SFDC number)."""
    _replace_body_text(doc, {
        CUSTOMER_PLACEHOLDER: meta.customer or "",
        QUOTE_NUMBER_PLACEHOLDER: meta.quote_number or "",
    })


def _fill_sales_contact(doc, meta: QuotationMeta) -> None:
    """Replace the sales-contact placeholders left in the template (if provided).

    Operates on every <w:t> in the body so it reaches the page-1 cover content
    control. The mailto link target is repointed to the provided email too.
    """
    from docx.oxml.ns import qn

    repl = {}
    if meta.sales_name:
        repl[SALES_NAME_PLACEHOLDER] = meta.sales_name
    if meta.sales_email:
        repl[SALES_EMAIL_PLACEHOLDER] = meta.sales_email
    if not repl:
        return
    for t in doc.element.body.iter(qn("w:t")):
        if not t.text:
            continue
        new = t.text
        for old, value in repl.items():
            if old in new:
                new = new.replace(old, value)
        if new != t.text:
            t.text = new
    if meta.sales_email:
        for rel in list(doc.part.rels.values()):
            try:
                target = rel.target_ref or ""
            except Exception:  # noqa: BLE001
                continue
            if "mailto:" in target and SALES_EMAIL_PLACEHOLDER in target:
                rel._target = f"mailto:{meta.sales_email}"


def _blank_all_prices(doc) -> None:
    """Blank every monetary amount in the body and tables (pending Step 3)."""
    for para in doc.paragraphs:
        if "$" in para.text:
            _set_text_keep_format(para, _blank_money_in_text(para.text))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "$" in para.text:
                        _set_text_keep_format(para, _blank_money_in_text(para.text))


# --------------------------------------------------------------------------- #
# Main builder
# --------------------------------------------------------------------------- #
def generate_quotation(
    step1: dict,
    step2: dict,
    output_path: str | Path,
    meta: Optional[QuotationMeta] = None,
    template_path: Optional[str | Path] = None,
) -> Path:
    """Clone the standard template and fill it from Step 1 + Step 2 outputs.

    Args:
        step1: parsed ``step1_extract_info.json``.
        step2: parsed ``step2_create_boq.json``.
        output_path: destination .docx path.
        meta: optional header metadata.
        template_path: override the standard quotation template.

    Returns:
        Path to the written document.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required: pip install python-docx") from exc

    from qualitrol_core import config

    project_id = step2.get("project_id") or step1.get("project_id") or "PROJECT"
    meta = meta or QuotationMeta(project_id=project_id)
    template = Path(template_path) if template_path else config.quotation_template_path()
    if not template.exists():
        raise FileNotFoundError(f"Quotation template not found: {template}")

    doc = Document(str(template))

    _fill_header(doc, meta)
    _fill_cover(doc, meta)
    _fill_revision(doc)
    _fill_project_info(doc, meta)
    _fill_customer_requirement(doc, step1)
    _fill_important_notes(doc, step2)
    _fill_equipment_schedule(doc, step1, step2)
    _fill_sales_contact(doc, meta)
    _blank_all_prices(doc)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out

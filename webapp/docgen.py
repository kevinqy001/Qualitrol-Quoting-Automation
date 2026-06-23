"""Word quotation document generation.

Ported from the original POC ``backend/app/core/docgen.py`` and extended to
accept the richer Step 1 + Step 2 output structures used in this project.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any


def generate_quotation_docx(priced_boq: dict[str, Any], output_path: str | Path) -> Path:
    """Generate a Word .docx quotation document from a priced BOQ.

    Args:
        priced_boq: Dict with keys: caseReference, currency, lineItems,
                    subtotal, tax, grandTotal, validityDays, paymentTerms.
                    lineItems entries need: productCode, description, quantity,
                    unitPrice, discountPercent, netUnitPrice, lineTotal.
        output_path: Where to save the .docx (directory is created if needed).

    Returns:
        Path to the generated file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx is required: pip install python-docx")

    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────── #
    title = doc.add_heading("Qualitrol — Quotation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Metadata ─────────────────────────────────────────────────────────── #
    meta = [
        ("Reference", priced_boq.get("caseReference", "—")),
        ("Currency", priced_boq.get("currency", "USD")),
        ("Date", datetime.now().strftime("%Y-%m-%d")),
        ("Validity", f"{priced_boq.get('validityDays', 90)} days"),
        ("Payment Terms", priced_boq.get("paymentTerms", "Net 30")),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    doc.add_paragraph()

    # ── Line Items ───────────────────────────────────────────────────────── #
    doc.add_heading("Bill of Quantities", level=1)
    items = priced_boq.get("lineItems", [])
    if items:
        col_headers = ["#", "Product Code", "Description", "Qty",
                       "Unit Price", "Disc %", "Net Unit", "Line Total"]
        table = doc.add_table(rows=1 + len(items), cols=len(col_headers))
        table.style = "Table Grid"

        hdr_row = table.rows[0]
        for i, header in enumerate(col_headers):
            cell = hdr_row.cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                run.bold = True

        def _fmt(val: Any, is_money: bool = False) -> str:
            if is_money:
                try:
                    return f"${float(val):,.2f}"
                except (TypeError, ValueError):
                    return str(val)
            return str(val) if val is not None else "—"

        for row_idx, item in enumerate(items, start=1):
            row = table.rows[row_idx]
            row.cells[0].text = _fmt(item.get("lineNumber", row_idx))
            row.cells[1].text = _fmt(item.get("productCode") or item.get("product_model", ""))
            row.cells[2].text = _fmt(item.get("description") or item.get("product_description", ""))
            row.cells[3].text = _fmt(item.get("quantity"))
            row.cells[4].text = _fmt(item.get("unitPrice", 0), is_money=True)
            row.cells[5].text = f"{item.get('discountPercent', 0)}%"
            row.cells[6].text = _fmt(item.get("netUnitPrice", 0), is_money=True)
            row.cells[7].text = _fmt(item.get("lineTotal", 0), is_money=True)
    else:
        doc.add_paragraph("No priced line items available.")

    # ── Totals ───────────────────────────────────────────────────────────── #
    doc.add_paragraph()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Subtotal: ${priced_boq.get('subtotal', 0):,.2f}")
    doc.add_paragraph(f"Tax: ${priced_boq.get('tax', 0):,.2f}")
    grand = doc.add_paragraph()
    run = grand.add_run(f"Grand Total: ${priced_boq.get('grandTotal', 0):,.2f}")
    run.bold = True
    run.font.size = Pt(14)

    # ── Missing Info section (if present) ────────────────────────────────── #
    missing = priced_boq.get("missingInfoQuestions", [])
    if missing:
        doc.add_page_break()
        doc.add_heading("Open Clarification Questions", level=1)
        doc.add_paragraph(
            "The following items require customer confirmation before the BOQ can be finalised."
        )
        for i, q in enumerate(missing, start=1):
            doc.add_heading(
                f"{i}. [{q.get('priority', 'Medium')}] {q.get('missing_item', 'Missing information')}",
                level=2,
            )
            if q.get("why_it_matters"):
                p = doc.add_paragraph()
                p.add_run("Why it matters: ").bold = True
                p.add_run(q["why_it_matters"])
            if q.get("question"):
                p = doc.add_paragraph()
                p.add_run("Question: ").bold = True
                p.add_run(q["question"])
            if q.get("owner"):
                doc.add_paragraph(f"Owner: {q['owner']}")

    # ── Footer ───────────────────────────────────────────────────────────── #
    doc.add_paragraph()
    doc.add_paragraph(
        f"Generated by Qualitrol Quotation Agent · {datetime.now().strftime('%Y-%m-%d %H:%M')} UTC"
    )

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
